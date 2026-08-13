import os
import pypdf
from docx import Document

class DocumentParser:
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract plain text from PDF file using pypdf with fallback."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extracted_text = []
        try:
            reader = pypdf.PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    extracted_text.append(page_text)
            
            full_text = "\n".join(extracted_text).strip()
            if full_text:
                return full_text
        except Exception:
            pass

        # Fallback for plain text files or raw text stream
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                fallback_text = f.read().strip()
                if fallback_text:
                    return fallback_text
        except Exception:
            pass

        raise ValueError("Unable to extract text from PDF document.")

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract plain text from DOCX file using python-docx with fallback."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        paragraphs.append(row_text)
                        
            full_text = "\n".join(paragraphs).strip()
            if full_text:
                return full_text
        except Exception:
            pass

        # Fallback for plain text files
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                fallback_text = f.read().strip()
                if fallback_text:
                    return fallback_text
        except Exception:
            pass

        raise ValueError("Unable to extract text from DOCX document.")

    @classmethod
    def parse_document(cls, file_path: str, file_type: str = None) -> str:
        """Extract text automatically based on file extension or mime type."""
        ext = os.path.splitext(file_path)[1].lower()

        try:
            if ext == '.pdf' or file_type == 'application/pdf':
                return cls.extract_text_from_pdf(file_path)
            elif ext in ['.docx', '.doc'] or 'wordprocessingml' in str(file_type):
                return cls.extract_text_from_docx(file_path)
            else:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read().strip()
        except Exception as e:
            # General fallback
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    txt = f.read().strip()
                    if txt:
                        return txt
            except Exception:
                pass
            raise ValueError(f"Document text extraction failed: {str(e)}")
